# -*- coding: UTF-8 -*-
import sys
import math
import re

from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from xml.dom.minidom import parseString
import xml.etree.ElementTree as ET
from zipfile import ZipFile, ZIP_DEFLATED
from lxml.etree import fromstring, tostring

from docx.opc.constants import RELATIONSHIP_TYPE as RT
reload(sys)
sys.setdefaultencoding('utf8')

def is_table_title(document, title):
    for p in document.paragraphs:
        if 'caption' in p.style.name.lower():
            if title in p.text.lower():
                return True
    return False


def get_cells(table):
    cells_values = []
    for cell in table._cells:
        text = ''
        for p in cell.paragraphs:
            if p.text.replace(" ", "") != '':
                text += p.text
        cells_values.append(text)
    return cells_values

def cells_align_center(table):
    for cell in table._cells:
        for p in cell.paragraphs:
            if p.text.replace(" ", "") != '':
                if not 'CENTER' in str(p.alignment):
                    return False
    return True


def string_arrays_is_equal(arr1, arr2):
    tmp_arr1 = []
    tmp_arr2 = []

    for tmp in arr1:
        tmp_arr1.append(tmp.replace(" ", ""))
    for tmp in arr2:
        tmp_arr2.append(tmp.replace(" ", ""))

    if (tmp_arr1 == tmp_arr2):
        return True
    else:
        return False

def lab_2_check_answer(student_path, correct_path):
    print '!!!!!!!!!!!!!!!!!!!!!!!!!!!!', student_path
    print '!!!!!!!!!!!!!!!!!!!!!!!!!!!!', correct_path
    response_message = {}
    response_message["errors"] = []
    response_message["table_format"] = {}
    response_message["table_title"] = {}
    response_message["cells_alignment"] = {}
    table_title = 'автобусные маршруты'
    try:
        document_correct = Document(correct_path)
        document_student = Document(student_path)
    except:
        response_message["errors"].append("При открытии документа произошла ошибка (документ должен иметь расширение docx)")
        return response_message

    if len(document_student.tables) == 1:
        table_student = document_student.tables[0]
        table_correct = document_correct.tables[0]
        # print table_student

        if string_arrays_is_equal(get_cells(table_student), get_cells(table_correct)):
            response_message["table_format"]["status"] = True
            response_message["table_format"]["message"] = "Форматирование таблицы выполнено верно"
        else:
            response_message["table_format"]["status"] = False
            response_message["table_format"]["message"] = "Форматирование таблицы выполнено неверно"

        if is_table_title(document_student, table_title):
            response_message["table_title"]["status"] = True
            response_message["table_title"]["message"] = "Заголовок таблицы выполнен верно"
        else:
            response_message["table_title"]["status"] = False
            response_message["table_title"]["message"] = "Заголовок таблицы выполнен неверно (не выполнен)"

        if cells_align_center(table_student):
            response_message["cells_alignment"]["status"] = True
            response_message["cells_alignment"]["message"] = "Выравнивание внутри ячеек (по центру) верно"
        else:
            response_message["cells_alignment"]["status"] = False
            response_message["cells_alignment"]["message"] = "Выравнивание внутри ячеек (по центру) неверно (не выполнен)"

        return response_message
    else:
        response_message["errors"].append("Документ должен содержать одну таблицу")
        return response_message