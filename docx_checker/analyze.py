import sys
import math
import re

from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from xml.dom.minidom import parseString
import xml.etree.ElementTree as ET
from lxml import etree

reload(sys)
sys.setdefaultencoding('utf8')

def get_analyze_the_document(path):
    file_analyzes = {}
    def get_general_properties(document):
        general_properties = {}
        general_properties["author"] = str(document.core_properties.author) # Автор документа
        general_properties["created"] = str(document.core_properties.created) # Дата создания документа
        general_properties["last_modified_by"] = str(document.core_properties.last_modified_by) # Пользователь, который менял последний
        general_properties["modified"] = str(document.core_properties.modified) # Время изменения
        general_properties["title"] = str(document.core_properties.title) #Название
        file_analyzes["general_properties"] = general_properties

    def get_document_margins(document):
        margins = {}
        sections = document.sections
        margins["top"] = str(round(sections[0].top_margin.cm, 2))
        margins["bottom"] = str(round(sections[0].bottom_margin.cm, 2))
        margins["left"] = str(round(sections[0].left_margin.cm, 2))
        margins["right"] = str(round(sections[0].right_margin.cm, 2))
        file_analyzes["margins"] = margins

    def get_headers_texts(document):
        document_headers_texts = []
        count = 0
        for p in document.paragraphs:
            if p.style.name != "Normal":
                if "Заголовок" in p.style.name and "ФИО" in p.style.name:
                    text = re.sub(r'\s+', ' ', p.text)
                    document_headers_texts.append((text))
            if 'toc' in p.style.name:
                count +=1

        file_analyzes["menu_item_count"] = count
        file_analyzes["document_headers_texts"] = document_headers_texts

    def get_subject_index(document):
        subject_index_temp = []
        subject_index = []
        fi = 0
        all_p = document.paragraphs
        for i, val in enumerate(all_p):
            if "Заголовок" in val.style.name and "предметный указатель" in val.text.lower():
                fi = i
        all_p = all_p[(fi+1):]

        for p in all_p:
            if "указатель" in p.style.name.lower():
                subject_index_temp.append(p.text)
        for s in subject_index_temp:
            i = s.rfind(',')
            subject_index.append(re.sub(r'\s+', ' ', s[:i]))
        file_analyzes["subject_index"] = subject_index

    def get_custom_styles(document):
        custom_styles = {}
        all_docx_styles = document.styles
        for s in all_docx_styles:
            if s.builtin == False and "ФИО" in s.name:
                if "Заголовок" in s.name:
                    header_style = {}
                    custom_head_style = all_docx_styles[s.base_style.name]
                    header_style["name"] = s.name
                    header_style["font_name"] = str(s.font.name)
                    header_style["font_size"] = str(s.font.size)
                    header_style["font_italic"] = str(s.font.italic)
                    header_style["font_bold"] = str(s.font.bold)
                    header_style["line_spacing"] = str(s.paragraph_format.line_spacing)
                    header_style["first_line_indent"] = str(s.paragraph_format.first_line_indent)
                    header_style["space_before"] = str(s.paragraph_format.space_before)
                    header_style["space_after"] = str(s.paragraph_format.space_after)
                    header_style["alignment"] = str(s.paragraph_format.alignment)
                    custom_styles["header_style"] = header_style
                else:
                    pass
                    # paragraph_style = {}
                    # custom_paragraph_style = all_docx_styles[s.base_style.name]
                    # header_style["name"] = s.name
                    # header_style["font_name"] = s.font.name

                    # print s.paragraph_format.alignment

        file_analyzes["custom_styles"] = custom_styles

    def get_document_numbering(document):
        def iter_header_parts(document):
            document_part = document.part
            for rel in document_part.rels.values():
                if rel.reltype == RT.FOOTER:
                    yield rel.target_part

        document_page_numbering = False
        for header_part in iter_header_parts(document):

            header_xml = header_part._blob
            namespace = dict(w="http://schemas.openxmlformats.org/wordprocessingml/2006/main")
            root = ET.fromstring(header_xml)
            # print header_xml
            try:
                text_element = root.find(".//w:docPartGallery", namespace)
                if "Page Numbers" in ET.tostring(text_element):
                    document_page_numbering = True
            except:
                pass
        file_analyzes["document_page_numbering"] = document_page_numbering

    def get_document_heading(document):
        document_header = ""
        def iter_header_parts(document):
            document_part = document.part
            for rel in document_part.rels.values():
                if rel.reltype == RT.HEADER:
                    yield rel.target_part


        for header_part in iter_header_parts(document):
            header_xml = header_part._blob
            namespace = dict(w="http://schemas.openxmlformats.org/wordprocessingml/2006/main")
            root = ET.fromstring(header_xml)
            text_element = root.find(".//w:t", namespace)
            if text_element is not None:
                if text_element.text != '':
                    document_header = text_element.text
        file_analyzes["document_header"] = document_header

    document = Document(path)
    get_headers_texts(document)
    get_general_properties(document)
    get_document_margins(document)
    get_custom_styles(document)
    get_subject_index(document)
    get_document_numbering(document)
    get_document_heading(document)
    return file_analyzes