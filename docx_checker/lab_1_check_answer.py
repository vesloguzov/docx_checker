# -*- coding: UTF-8 -*-
import sys
import math
import re

from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from xml.dom.minidom import parseString
import xml.etree.ElementTree as ET
from zipfile import ZipFile, ZIP_DEFLATED
from lxml.etree import fromstring, tostring

from docx.opc.constants import RELATIONSHIP_TYPE as RT
reload(sys)
sys.setdefaultencoding('utf8')



def get_custom_header_style(document, header_name):
    custom_head_style = {}
    for s in document.styles:
        if s.builtin == False and header_name in s.name:
            custom_head_style["base_style"] = s.base_style.name
            custom_head_style["font_name"] = s.font.name
            custom_head_style["font_italic"] = s.font.italic
            custom_head_style["font_bold"] = s.font.bold
            custom_head_style["space_before"] = s.paragraph_format.space_before.pt
            custom_head_style["space_after"] = s.paragraph_format.space_after.pt
            custom_head_style["alignment"] = s.paragraph_format.alignment
    return custom_head_style

def get_custom_main_style(document, main_name):
    custom_main_style = {}
    for s in document.styles:
        if s.builtin == False and main_name in s.name:
            custom_main_style["base_style"] = s.base_style.name
            custom_main_style["font_name"] = s.font.name
            custom_main_style["font_italic"] = s.font.italic
            custom_main_style["font_bold"] = s.font.bold
            custom_main_style["line_spacing"] = s.paragraph_format.line_spacing
            custom_main_style["first_line_indent"] = s.paragraph_format.first_line_indent.cm
            custom_main_style["space_before"] = s.paragraph_format.space_before.pt
            custom_main_style["space_after"] = s.paragraph_format.space_after.pt
            custom_main_style["alignment"] = s.paragraph_format.alignment
    return custom_main_style

def get_document_margins(document):
    document_margins = {}
    sections = document.sections
    try:
        document_margins["top"] = str(round(sections[0].top_margin.cm, 2))
        document_margins["bottom"] = str(round(sections[0].bottom_margin.cm, 2))
        document_margins["left"] = str(round(sections[0].left_margin.cm, 2))
        document_margins["right"] = str(round(sections[0].right_margin.cm, 2))
    except:
        pass
    return document_margins


def get_document_header(document):
    document_header = ""
    try:
        def iter_header_parts(document):
            document_part = document.part
            for rel in document_part.rels.values():
                if rel.reltype == RT.HEADER:
                    yield rel.target_part
        for header_part in iter_header_parts(document):
            header_xml = header_part._blob
            namespace = dict(w="http://schemas.openxmlformats.org/wordprocessingml/2006/main")
            root = ET.fromstring(header_xml)
            text_elements = root.findall(".//w:t", namespace)
            if text_elements is not None:
                for element in text_elements:
                    document_header += element.text
    except:
        pass
    
    return document_header

def get_footnotes(document_path):
    footnotes = []
    try:
        sourceFile = ZipFile(document_path, 'r')
        namespace = dict(w="http://schemas.openxmlformats.org/wordprocessingml/2006/main")
        charts = []
        [charts.append(sourceFile.read(ch)) for ch in sourceFile.namelist() if 'footnotes' in ch]
        chart = ET.fromstring(charts[0])
        footnotes_list = chart.findall(".//w:footnote", namespace)
        for footnote in footnotes_list:
            current_footnote = ""
            text_elements = footnote.findall(".//w:t", namespace)
            if text_elements is not None:
                for element in text_elements:
                    current_footnote += element.text
            if(current_footnote.replace(" ", "") != ''):
                footnotes.append(current_footnote)
    except:
        pass
    return footnotes

def is_document_numbering(document):
    def iter_header_parts(document):
        document_part = document.part
        for rel in document_part.rels.values():
            if rel.reltype == RT.FOOTER:
                yield rel.target_part

    document_page_numbering = False
    for header_part in iter_header_parts(document):

        header_xml = header_part._blob
        namespace = dict(w="http://schemas.openxmlformats.org/wordprocessingml/2006/main")
        root = ET.fromstring(header_xml)
        # print header_xml
        try:
            text_element = root.find(".//w:docPartGallery", namespace)
            if "Page Numbers" in ET.tostring(text_element):
                document_page_numbering = True
        except:
            pass
    return document_page_numbering

def is_table_of_contents(document):
    # Доделать
    for p in document.paragraphs:
        if 'toc' in p.style.name:
            return True
    return False

def get_docement_headers(document, header_name):
    headers = []
    for p in document.paragraphs:
        if header_name in p.style.name:
            if p.text != '':
                text = re.sub(r'\s+', ' ', p.text)
                headers.append(text)
    return len(headers)


def lab_1_check_answer(student_path, correct_path):
    response_message = {}
    response_message["errors"] = []
    response_message["custom_header_style"] = {}
    response_message["custom_main_style"] = {}
    response_message["document_margins"] = {}
    response_message["document_header"] = {}
    response_message["footnote"] = {}
    response_message["table_of_contents"] = {}
    response_message["document_numbering"] = {}
    response_message["headers_style"] = {}
    try:
    # document_path = 'data/internet_tables.docx'
        document_student = Document(student_path)
        document_correct = Document(correct_path)
        
        header_style_name = "Заголовок_ФИО"
        main_style_name = "Основной_ФИО"

        if is_table_of_contents(document_student):
            response_message["table_of_contents"]["message"] = "Оглавление создано"
            response_message["table_of_contents"]["status"] = True
        else:
            response_message["table_of_contents"]["message"] = "Оглавление не создано"
            response_message["table_of_contents"]["status"] = False

        if is_document_numbering(document_student):
            response_message["document_numbering"]["message"] = "Страницы пронумерованы"
            response_message["document_numbering"]["status"] = True
        else:
            response_message["document_numbering"]["message"] = "Страницы не пронумерованы"
            response_message["document_numbering"]["status"] = False

        main_style_student = get_custom_header_style(document_student, main_style_name)
        main_style_correct = get_custom_header_style(document_correct, main_style_name)

        if main_style_student == main_style_correct:
            response_message["custom_main_style"]["message"] = "Основной стиль создан верно"
            response_message["custom_main_style"]["status"] = True
        else:
            response_message["custom_main_style"]["message"] = "Основной стиль создан неверно"
            response_message["custom_main_style"]["status"] = False

        header_style_student = get_custom_header_style(document_student, header_style_name)
        header_style_correct = get_custom_header_style(document_correct, header_style_name)

        if header_style_student == header_style_correct:
            response_message["custom_header_style"]["message"] = "Стиль заголовков создан верно"
            response_message["custom_header_style"]["status"] = True
        else:
            response_message["custom_header_style"]["message"] = "Стиль заголовков создан неверно"
            response_message["custom_header_style"]["status"] = False

        if get_footnotes(student_path) == get_footnotes(correct_path):
            response_message["footnote"]["message"] = "Сноска создана верно"
            response_message["footnote"]["status"] = True
        else:
            response_message["footnote"]["message"] = "Сноска создана неверно"
            response_message["footnote"]["status"] = False
       
        if get_docement_headers(document_student, header_style_name) == get_docement_headers(document_correct, header_style_name):
            response_message["headers_style"]["message"] = "Стиль к заголовкам применен верно"
            response_message["headers_style"]["status"] = True
        else:
            response_message["headers_style"]["message"] = "Стиль к заголовкам применен неверно"
            response_message["headers_style"]["status"] = False

        if get_document_header(document_student) == get_document_header(document_correct):
            response_message["document_header"]["message"] = "Верхний колонтитул применен верно"
            response_message["document_header"]["status"] = True
        else:
            response_message["document_header"]["message"] = "Верхний колонтитул применен неверно"
            response_message["document_header"]["status"] = False


        if get_document_margins(document_student) == get_document_margins(document_correct):
            response_message["document_margins"]["message"] = "Отступы применены верно"
            response_message["document_margins"]["status"] = True
        else:
            response_message["document_margins"]["message"] = "Отступы применены неверно"
            response_message["document_margins"]["status"] = False
    except:
        response_message["errors"].append('Ошибка при открытии файла')
    return response_message
# is_document_numbering(document)
